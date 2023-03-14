from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_resume_data():
    first_name = "Daniel"
    last_name = "Fredriksen"

    address = "Las Vegas, Honolulu, New York"
    phone = "+1-702-408-4578"
    email = "dfredriksen@cyint.technology"
    linkedin = "https://linkedin.com/in/cyint"
    calendly = "https://calendly.com/d/d68-yvd-bjh"

    titles = [
        "Machine Learning Engineer",
        "Data Scientist",
        "Software Engineer",
        "NLP Engineer",
        "NLU Engineer",
        "Prompt Engineer",
        "Full Stack Developer",
        "Front-end Developer",
        "UI/UX Developer",
        "React Developer",
        "Software Architect",
        "Quantum Computing Engineer",
        "Quantum Computing Scientist",
        "Machine Learning Scientist"
    ]

    headline = "Expert Machine Learning Engineer / Data Scientist"
    summary = "I am an expert machine learning engineer and data scientist having over 26 years of experience writing software. I've spent the last 10 years building data pipelines, training machine learning models, and productionizing insights. My track record shows a pattern of delivering value."

    technical_skills = {
        "Data Science": [
            "Cleaning",
            "Wrangling",
            "Tidying",
            "Juypter Notebooks",
            "Sentiment Analysis",
            "Switchpoint Analysis",
            "Time Series Analysis",
            "Data Analytics",
            "Statistical Testing",
            "Experimental Design",
            "A/B Testing",
            "Data Visualization",
            "Reproduceable Research",
            "Cybersecurity",
            "Operations Science",
            "Pareto Analysis",
            "Anomaly Detection",
            "Conformal Prediction"
        ],
        "Data Engineering": [
            "ETL",
            "ELT",
            "SQL",
            "AWS Glue",
            "Azure Data Factory",
            "Databricks",
            "Snowflake",
            "Postgres",
            "Mongo",
            "Cosmos",
            "Redshift",
            "Aurora",
            "Dynamo",
            "EMR",
            "Hadoop",
            "Big Query",
            "Airflow",
            "MS SQL Server"
        ],
        "Software Engineering": [
            "Python",
            "Node",
            "C#",
            "C/C++",
            "Go",
            "PHP",
            "React",
            "Lambda",
            "Azure Functions",
            "Git / GitHub",
            "GitHub Actions",
            "Bitbucket Actions",
            "Azure DevOps",
            "Docker",
            "Kubernetes",
            "Jenkins",
            "Travis CI",
            "Team City",
            "Fargate",
            "CloudFront",
            "S3",
            "EC2",
            "VPC",
            "SAS",
            "Java",
            "DevOps",
            "CICD",
            "Unit Testing",
            "End to End Testing",
            "Selenium",
            "PyTest",
            "Jest",
            "Terraform",
            "Serverless"
        ],
        "Machine Learning Engineering": [
            "Linear Regression",
            "Logistic Regression",
            "Decision Trees",
            "Random Forest",
            "Convolutional Neural Nets (CNN)",
            "GPT-3 / GPT-3.5 Turbo",
            "Clustering",
            "Personalization",
            "Supervised",
            "Unsupervised",
            "PyTorch",
            "Tensorflow",
            "TensorflowJS",
            "Tensorflow Lite",
            "Hugging Face",
            "LSTM",
            "Genetic Algorithms",
            "XGBoost",
            "Sklearn",
            "Reinforcement Learning",
            "Named Entity Recognition",
            "BERT",
            "TorchServe",
            "Vertex AI",
            "Sagemaker",
            "Azure ML"
        ],
        "Cloud Computing": [
            "AWS Professional Certification",
            "GCP Professional Certification",
            "Azure Professional Certification"
        ],
        "Project / Product Management": [
            "JIRA",
            "Business Analysis",
            "Technical Writing",
            "Agile",
            "SCRUM",
            "Asyncronous Work",
            "Sprint planning",
            "Backlog grooming",
            "UI/UX",
            "Technical Product Management",
            "OKRs"
        ],
        "Business Intelligence / Analytics": [
            "Dask",
            "Plotly",
            "Tableau",
            "Power BI",
            "Looker",
            "Google Analytics",
            "Google Tag Manager",
            "Segment IO",
            "Mixpanel",
        ],
        "Crypto": [
            "Solidity",
            "Dapps",
            "IPFS",
            "Blockchain"
        ],
        "Quantum Computing": [
            "Qiskit",
            "Quantum Support Vector Machines"
        ],
    }

    work_history = [
    {
        "company": "CYINT LLC",
        "title": "Chief Data Scientist",
        "dates": "4/2007 - Present",
        "description": "Consulting for various data science and machine learning projects, including high profile customers such as the NYC Office of the Chief Medical Examiner, the Veteran's Health Administration, Toyota, Levis, DocuSign, CVS Health, TC Energy, Experian, Expedia, and others",
        "skills": [technical_skills[skill] for skill in technical_skills.keys()],
        "results": [
            "Productionized ML Models",
            "Deployed software using cloud services",
            "Conducted R&D",
            "Developed ELT and ETL Pipelines",
            "Provided mentorship and guidance to junior team members resulting in productivity gains",
            "Developed front-end user interfaces",
            "Implemented A/B testing frameworks",
            "Developed dashboards",
            "Improved team velocity",
            "Estimated project efforts",
            "Curated backlog",
            "Managed Agile teams",
            "Produced analytical insights using data science"
        ],
    }, {
        "company":"Enterprise Technology Research",
        "title": "Chief Data Scientist",
        "dates": "7/2020 - 12/2021",
        "description": "Developed and productionized machine learning models to provide insights about companies in the enterprise technology space to investors, launching their first forecasting models and reducing customer churn by 10%",
        "skills": [
            "Data Science",
            "Python",
            "MS SQL Server",
            "SAS",
            "Power BI",
            "Time Series Analysis",
            "Analytical Reporting",
            "JIRA",
            "Agile",
            "Asyncronous Work",
            "Selenium",
            "Survey Design",
            "GPT-3",
            "Sentiment Analysis",
            "NLP",
            "Entity extraction",
            "A/B Testing",
            "React"
        ],
        "results": [
            "Forecasted the largest Microsoft earnings call consensus beat in years",
            "Reduced customer churn by 10% through delivery of forecasting models",
            "Developed and deployed software to automatically generate intelligence report summaries",
            "Automated intelligence reports based on data collected leveraging GPT-3, increasing efficiency by 40x",
            "Implemented ELT pipelines to capture and collect data powering $7M in revenues",
            "Managed teams of data analysts, data scientists, and ML engineers to drive OKRs"
        ],
    }, {
        "company":"ESW Capital (Crossover, ZephyrTel)",
        "title": "VP of Engineering and Operations",
        "dates": "6/2018 - 1/2020",
        "description": "Led software engineering efforts and business operations for a large private equity firm. \
            Transformed aquisition targets to achieve a 70% operating margin within 90 days through shift & lift\
            manuevers into the cloud, software redesign, and the application of data science for operations research",
        "skills": [
            "Data Science",
            "Python",
            "C/C++",
            "Java",
            "Hadoop",
            "EC2",
            "Aurora",
            "Node",
            "Team City",
            "Jenkins",
            "Technical Product Management"
            "Operations Research",
            "Software Engineering",
            "Machine Learning Engineering",
            "Technical Support",
            "JIRA",
            "Agile",
            "Asyncronous Work",
            "Selenium",
            "Sentiment Analysis",
            "NLP",
            "Entity extraction",
            "Chatbot",
            "A/B Testing",
            "React"
        ],
        "results": [
            "Transformed three acquisitions to achieve 70% operating margin within 90 days",
            "Oversaw engineering and operations for multi-million dollar assets in the telecom and ML space",
            "Extracted product insights from support data to identify and eliminate the root causes of product dissatisfaction",
            "Automated technical support by implementing chatbots powered by NLU and NLP models",
            "Surgical modification of software that resulted in million dollar efficiency gains",
            "Managed a large number of direct reports and dotted line reports across different departments and business units"
        ],
    }, {
        "company":"Coolmath Games",
        "title": "VP of Technology Systems Operations",
        "dates": "1/2016 - 12/2016",
        "description": "A/B tested different header bidding services resulting in a revenue boost of 30%, \
            Implemented Stripe integration into the platform to enable subscriptions, \
            Developed a web traffic anomaly model to reduce downtime from DOS attacks to zero, \
            Managed a team of developers using agile practices, \
            Maintained a Drupal system hosted on Acquia to service over 6M monthly unique users",
        "skills": [
            "Data Science",
            "PHP",
            "Drupal",
            "Acquia",
            "Header Bidding",
            "A/B Testing",
            "Javascript",
            "Jenkins",
            "Technical Product Management"
            "Machine Learning",
            "Anomaly Detection",
            "Nginx",
            "JIRA",
            "Agile",
        ],
        "results": [
            "A/B tested different header bidding services resulting in a revenue boost of 30%",
            "Implemented Stripe integration into the platform to enable subscriptions",
            "Developed a web traffic anomaly model to reduce downtime from DOS attacks to zero",
            "Managed a team of developers using agile practices",
            "Maintained a Drupal system hosted on Acquia to service over 6M monthly unique users",
        ],
    },{
        "company":"Stealth Travel",
        "title": "Senior Front-End Architect",
        "dates": "6/2015 - 12/2016",
        "description": "Led the front-end team for a React based travel startup, \
            delivering sprint goals at 100% on-time and developing an innovative solution \
            that leveraged the V8 engine to optionally render components server-side. The startup \
            closed down after investors pulled out due to unforseen circumstances",
        "skills": [
            "React",
            "PHP",
            "Docker",
            "Styleguides",
        ],
        "results": [
            "Developed React components for a vivid travel website with pixel-perfect accuracy",
            "Implemented a styleguide approach resulting in modular component delivery",
            "Innovated a new server-side rendering approach using a V8/PHP integration eliminating \
                 the need to re-develop components in PHP",
        ],
    }, {
        "company":"DNA Info",
        "title": "Full Stack Software Engineer / Data Scientist",
        "dates": "5/2012 - 6/2015",
        "description": "Developed and maintained a website for a digital news agency with over 3M monthly \
            unique visitors. Introduced CICD best practices and Data Science for A/B testing, personalization \
            to drive key metrics such as engagement and time-on-site",
        "skills": [
            "PHP",
            "Python",
            "Docker",
            "Nginx",
            "EC2",
            "Data Science",
            "Personalization",
            "Segment IO",
            "Sentiment Analysis",
            "Bamboo",
            "Cybersecurity",
            "DynamoDB",
        ],
        "results": [
            "Implemented a CICD pipeline using bamboo to reduce downtime due to deployments to zero",
            "Developed an A/B testing framework to test features, newsletter variants",
            "Introduced a personalization model based on cluster analysis",
        ]
    }
    ]

    education = [{
        "degree": "PhD in Quantum Computing",
        "university": "Capitol Technology University",
        "dates": "2021 - Present",
    }, {
        "degree": "Masters in Cybersecurity",
        "university": "Excelsior College",
        "dates": "2015 (Incomplete)",
    }, {
        "degree": "Masters in Intelligence Management",
        "university": "Henley Putnam University",
        "dates": "2015 (Incomplete)",
    }, {
        "degree": "Bachelors in Computer Science",
        "university": "Stonybrook University",
        "dates": "2011 (Incomplete)",
    }, {
        "degree": "Bachelors in Intelligence Management",
        "university": "Henley Putnam University",
        "dates": "2011",
    }, {
        "degree": "Associate in Science, Computer Science",
        "university": "Excelsior College",
        "dates": "2009",
    }]

    certifications = f"Over 70+ certifications in data science, machine learning, and software engineering \
     available on Linked In at: {linkedin}"

    return [
        first_name, 
        last_name, 
        address, 
        phone,
        email,
        calendly,
        linkedin,
        titles,
        headline,
        summary,
        technical_skills,
        work_history,
        education,
        certifications
    ]


def build_resume(
    cover_letter,
    first_name, 
    last_name, 
    address, 
    phone, 
    email, 
    calendly,
    linkedin, 
    title, 
    headline, 
    summary, 
    technical_skills, 
    work_history,
    education,
    certifications,
    filename
):
    
    document = Document()

    document.add_paragraph(cover_letter)
    document.add_page_break()

    document.add_heading(f"{first_name} {last_name}", 0)
    
    p = document.add_heading(headline, level=1)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"""{email}
{phone}
{calendly} 
{address}""")
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(summary)
    
    document.add_heading("Skills", level=1)
    for skill in technical_skills.keys():
        p = document.add_paragraph('')
        p.add_run(f"{skill}: ").bold = True
        p.add_run(", ".join(technical_skills[skill]))

    document.add_heading("Work History (Last 10 years)", level=1)

    for work in work_history:
        p = document.add_paragraph('')
        p.add_run({work['title']}).bold = True
        p.add_run(f"""
{work['company']}""")
        p.add_run(f"""
{work['dates']}""").italic = True
        p.add_run(work['description'])
        
    document.add_heading("Education", level=1)

    for degree in education:
        p = document.add_paragraph('')
        p.add_run(degree['degree']).bold = True
        p.add_run(f"""
{degree['university']}""")
        p.add_run(f"""
{degree['dates']}""").italic = True

    document.add_heading("Certifications", level=1)
    document.add_paragraph(certifications)
    document.add_paragraph('')

    document.add_heading("Military Service", level=1)
    document.add_paragraph("Served in support of Operation Iraqi Freedom III, and Operation Enduring Freedom VII \
        as an infantryman. Honorable discharge, receiving several medals and the rank of Sergeant")

    document.save(filename)
    return document
